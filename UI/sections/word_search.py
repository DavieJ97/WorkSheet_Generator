import random
import string
from docx.enum.text import WD_ALIGN_PARAGRAPH
from UI.exports.formatting import add_text
from UI.utils.editable_text import EditableText
from PyQt6.QtWidgets import (
    QFrame, QLabel,  QVBoxLayout,
    QHBoxLayout, QGridLayout, QPushButton,
    QComboBox,
)
import copy

class WordSearchSection(QFrame):

    def __init__(self, parent=None):
        super().__init__()

        self.main_window = parent
        self.word_inputs = []
        self.setFrameShape(QFrame.Shape.StyledPanel)
        self.setStyleSheet("""
            QFrame {
                background: white;
                border: 1px solid #d0d0d0;
                border-radius: 8px;
                padding: 10px;
            }
        """)

        self.init_ui()

    def init_ui(self):

        layout = QVBoxLayout()
        layout.setSpacing(8)

        # Section control buttons
        button_layout = QHBoxLayout()
        button_layout.addStretch()

        move_up = QPushButton("↑")
        move_down = QPushButton("↓")
        delete_button = QPushButton("Delete")

        move_up.clicked.connect(self.move_up)
        move_down.clicked.connect(self.move_down)
        delete_button.clicked.connect(self.delete_section)

        button_layout.addWidget(move_up)
        button_layout.addWidget(move_down)
        button_layout.addWidget(delete_button)

        layout.addLayout(button_layout)

        # Title
        self.title_input = EditableText(self.main_window, "Find the words")
        layout.addWidget(self.title_input)

        # Difficulty
        difficulty_layout = QHBoxLayout()
        difficulty_layout.addWidget(QLabel("Difficulty:"))
        self.difficulty = QComboBox()
        self.difficulty.addItems(["Beginner", "Medium", "Hard"])
        self.difficulty.currentTextChanged.connect(self.update_instruction_placeholder)
        difficulty_layout.addWidget(self.difficulty)
        layout.addLayout(difficulty_layout)

        # Grid size
        size_layout = QHBoxLayout()
        size_layout.addWidget(QLabel("Grid Size:"))
        self.grid_size = QComboBox()
        self.grid_size.addItems([
            "Small (8 x 8)",
            "Medium (12 x 12)",
            "Large (15 x 15)",
            "Extra Large (20 x 20)"
        ])
        self.grid_size.currentTextChanged.connect(self.update_word_limits)
        size_layout.addWidget(self.grid_size)
        layout.addLayout(size_layout)

        # Instructions
        instructions_label = QLabel("Instructions:")
        instructions_label.setStyleSheet("font-size:16px; font-weight:bold;")
        instructions_label.setFixedHeight(40)
        layout.addWidget(instructions_label)

        self.instructions = EditableText(self.main_window)
        layout.addWidget(self.instructions)

        # Word grid
        words_label = QLabel("Words:")
        words_label.setStyleSheet("font-size:16px; font-weight:bold;")
        words_label.setFixedHeight(40)
        layout.addWidget(words_label)

        grid = QGridLayout()

        rows = 3
        cols = 3

        for r in range(rows):
            for c in range(cols):

                field = EditableText(self.main_window)
                field.setMaxLength(12)
                field.textChanged.connect(lambda text, f=field: f.setText(text.upper()))

                self.word_inputs.append(field)

                grid.addWidget(field, r, c)

        layout.addLayout(grid)

        self.setLayout(layout)

        self.update_instruction_placeholder(self.difficulty.currentText())

    def update_instruction_placeholder(self, difficulty):
        if difficulty == "Beginner":
            self.instructions.setText(
                "The words can be ➡⬇"
            )
        elif difficulty == "Medium":
            self.instructions.setText(
                "The words can be ➡⬇↗↘"
            )
        elif difficulty == "Hard":
            self.instructions.setText(
                "The words can be ➡⬅⬆⬇↗↘↙↖"
            )

    def get_words(self):
        words = []
        for field in self.word_inputs:
            word = field.text().strip().upper()
            if word:
                words.append(word)

        return words
    
    def get_directions(self):
        difficulty = self.difficulty.currentText()
        if difficulty == "Beginner":
            return [(1,0), (0,1)]  # down, right
        elif difficulty == "Medium":
            return [(1,0), (0,1), (1,1), (-1,1)]
        else:
            return [
                (1,0), (0,1), (-1,0), (0,-1),
                (1,1), (-1,1), (1,-1), (-1,-1)
            ]
        
    def create_grid(self):
        size = self.get_grid_size()
        grid = [["" for _ in range(size)] for _ in range(size)]
        return grid
    
    def get_grid_size(self):
        size_map = {
            "Small (8 x 8)": 8,
            "Medium (12 x 12)": 12,
            "Large (15 x 15)": 15,
            "Extra Large (20 x 20)": 20
        }
        return size_map[self.grid_size.currentText()]
    
    def place_word(self, grid, word):
        size = len(grid)
        directions = self.get_directions()

        for attempt in range(100):
            direction = random.choice(directions)
            dr, dc = direction
            row = random.randint(0, size-1)
            col = random.randint(0, size-1)
            end_row = row + dr*(len(word)-1)
            end_col = col + dc*(len(word)-1)
            if not (0 <= end_row < size and 0 <= end_col < size):
                continue

            fits = True
            for i in range(len(word)):
                r = row + dr*i
                c = col + dc*i
                if grid[r][c] not in ("", word[i]):
                    fits = False
                    break
            if fits:
                for i in range(len(word)):
                    r = row + dr*i
                    c = col + dc*i
                    grid[r][c] = word[i]
                return True
        return False
    
    def fill_empty(self, grid):
        for r in range(len(grid)):
            for c in range(len(grid)):

                if grid[r][c] == "":
                    grid[r][c] = random.choice(string.ascii_uppercase)

    def generate_puzzle(self):
        grid = self.create_grid()
        size = self.get_grid_size()
        words = [w for w in self.get_words() if len(w) <= size]
        if not words:
            return self.create_grid()
        
        for word in words:
            placed = self.place_word(grid, word)
            if not placed:
                print(f"Could not place word: {word}")
        self.solution_grid = copy.deepcopy(grid)
        self.fill_empty(grid)
        return grid
    
    def get_size(self):
        grid = self.generate_puzzle()
        return len(grid) + 5
    
    def export(self, doc):
        instructions = self.instructions.text()
        if instructions:
            add_text(doc, instructions, italic=True, size=12)

        grid = self.generate_puzzle()
        table = doc.add_table(rows=len(grid), cols=len(grid))
        for r in range(len(grid)):
            for c in range(len(grid)):
                cell = table.rows[r].cells[c]
                cell.text = grid[r][c]
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        words = sorted(self.get_words())
        if words:
            add_text(doc, "Words", bold=True, size=12)
            word_line = "   ".join(words)
            add_text(doc, word_line, size=12)

    def export_answers(self, doc):
        doc.add_paragraph("Word Search Answer Key")
        grid = getattr(self, "solution_grid", None)
        if not grid:
            grid = self.generate_puzzle()
            grid = self.solution_grid

        table = doc.add_table(rows=len(grid), cols=len(grid))
        for r in range(len(grid)):
            for c in range(len(grid)):
                letter = grid[r][c]
                if letter:
                    table.rows[r].cells[c].text = letter
                else:
                    table.rows[r].cells[c].text = "."

    def delete_section(self):
        if self in self.main_window.sections:
            self.main_window.sections.remove(self)

        self.setParent(None)

    def move_up(self):
        layout = self.main_window.section_layout
        index = layout.indexOf(self)

        if index > 0:
            # Move widget
            layout.removeWidget(self)
            layout.insertWidget(index - 1, self)
            # Move section in list
            sections = self.main_window.sections
            sections[index], sections[index - 1] = sections[index - 1], sections[index]

    def move_down(self):
        layout = self.main_window.section_layout
        index = layout.indexOf(self)
        if index < layout.count() - 1:
            layout.removeWidget(self)
            layout.insertWidget(index + 1, self)
            sections = self.main_window.sections
            sections[index], sections[index + 1] = sections[index + 1], sections[index]

    def get_title(self):
        return self.title_input.text()
    
    def get_max_word_length(self):
        return self.get_grid_size()
    
    def update_word_limits(self):
        max_length = self.get_max_word_length()

        for field in self.word_inputs:
            field.setMaxLength(max_length)

            def handle_text_change(text, f=field):
                if len(text) >= max_length:
                    f.setToolTip(f"Cannot add words longer than {max_length} letters")
                else:
                    f.setToolTip("")

            field.textChanged.connect(handle_text_change)