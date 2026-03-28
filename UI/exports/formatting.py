from docx.shared import Pt

def add_formatted_text(doc, widget):

    text = widget.text()

    if not text:
        return

    font = widget.font()

    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)

    run.bold = font.bold()
    run.italic = font.italic()
    run.underline = font.underline()

    # Safe font size handling
    if font.pointSize() > 0:
        run.font.size = Pt(font.pointSize())
    else:
        run.font.size = Pt(14)  # default fallback

def add_text(doc, text, bold=False, italic=False, underline=False, size=None):
    if not text:
        return

    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)

    run.bold = bold
    run.italic = italic
    run.underline = underline

    # Only set size if explicitly provided
    if size is not None:
        run.font.size = Pt(size)

    return paragraph