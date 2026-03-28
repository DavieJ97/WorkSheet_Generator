from docx import Document
from UI.exports.formatting import add_formatted_text
from docx.shared import Pt


def export_document(title, sections, fill_path, headers=None):
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Worksheet title
    doc.add_heading(title, level=0)

    if headers:
        table = doc.add_table(rows=1, cols=len(headers)*2)

        col_index = 0
        for header in headers:
            table.rows[0].cells[col_index].text = f"{header}:"
            table.rows[0].cells[col_index + 1].text = "__________"
            col_index += 2

    # Page size system
    PAGE_CAPACITY = 33   # you can tweak this later
    current_size = 0

    for section in sections:
        section_title_widget = section.title_input

        # Get estimated size
        section_size = section.get_size()

        # 🔥 Check if it fits
        if current_size + section_size > PAGE_CAPACITY:
            doc.add_page_break()
            current_size = 0

        # Add section title (user-controlled formatting)
        add_formatted_text(doc, section_title_widget)

        # Add section content
        section.export(doc)

        # Update used space
        current_size += section_size

    doc.save(fill_path)

def export_answers(title, sections, file_path):
    doc = Document()
    doc.add_heading(f"{title} - Answer Key", level=0)
    for section in sections:
        section_title = section.get_title()
        doc.add_heading(section_title, level=1)
        section.export_answers(doc)
    doc.save(file_path)